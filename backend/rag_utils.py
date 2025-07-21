import os
import json
import pickle
from typing import List, Dict, Tuple, Optional, Any
from io import BytesIO
import logging

import PyPDF2
from docx import Document
import requests
import faiss
import numpy as np
from sentence_transformers import SentenceTransformer

from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document as LangchainDocument


logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

class RAGManager:
    """
    Manages RAG (Retrieval-Augmented Generation) functionality for the file sharing app.
    Handles text extraction, chunking, embeddings, and vector search.
    """
    
    def __init__(self, vector_db_path: str = "backend/vector_db", embedding_model: str = "all-MiniLM-L6-v2"):
        """
        Initialize RAG Manager
        
        Args:
            vector_db_path: Path to store vector databases
            embedding_model: Sentence transformer model name
        """
        self.vector_db_path = vector_db_path
        self.embedding_model_name = embedding_model
        self.embedding_model = None
        self.embedding_dimension = 384
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=800,
            chunk_overlap=100,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""]
        )
        
        os.makedirs(self.vector_db_path, exist_ok=True)
        
        self._initialize_embedding_model()
    
    def _initialize_embedding_model(self):
        """Initialize the sentence transformer model for embeddings"""
        try:
            logger.info(f"Loading embedding model: {self.embedding_model_name}")
            self.embedding_model = SentenceTransformer(self.embedding_model_name)
            dummy_embedding = self.embedding_model.encode(["test"])
            self.embedding_dimension = dummy_embedding.shape[1]
            logger.info(f"Embedding model loaded. Dimension: {self.embedding_dimension}")
        except Exception as e:
            logger.error(f"Failed to load embedding model: {e}")
            raise
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> str:
        """
        Extract text from various file formats
        
        Args:
            file_content: Raw bytes of the file
            filename: Name of the file (used to determine format)
            
        Returns:
            Extracted text content
        """
        file_extension = filename.lower().split('.')[-1]
        
        try:
            if file_extension == 'pdf':
                return self._extract_from_pdf(file_content)
            elif file_extension == 'docx':
                return self._extract_from_docx(file_content)
            elif file_extension == 'txt':
                return self._extract_from_txt(file_content)
            else:
                logger.warning(f"Unsupported file type: {file_extension}")
                return ""
        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            return ""
    
    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        text = ""
        try:
            pdf_reader = PyPDF2.PdfReader(BytesIO(file_content))
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
        return text.strip()
    
    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        text = ""
        try:
            doc = Document(BytesIO(file_content))
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
        return text.strip()
    
    def _extract_from_txt(self, file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            return file_content.decode('utf-8', errors='ignore')
        except Exception as e:
            logger.error(f"TXT extraction error: {e}")
            return ""
    
    def chunk_text(self, text: str, metadata: Dict) -> List[Dict]:
        """
        Split text into chunks with metadata
        
        Args:
            text: Text to chunk
            metadata: Metadata to attach to each chunk
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if not text.strip():
            return []
        
        doc = LangchainDocument(page_content=text, metadata=metadata)
        
        chunks = self.text_splitter.split_documents([doc])
        
        chunk_dicts = []
        for i, chunk in enumerate(chunks):
            chunk_dict = {
                'text': chunk.page_content,
                'metadata': {
                    **chunk.metadata,
                    'chunk_index': i,
                    'chunk_id': f"{metadata.get('filename', 'unknown')}_{i}"
                }
            }
            chunk_dicts.append(chunk_dict)
        
        return chunk_dicts
    
    def generate_embeddings(self, texts: List[str]) -> np.ndarray:
        """
        Generate embeddings for a list of texts
        
        Args:
            texts: List of text chunks
            
        Returns:
            Numpy array of embeddings
        """
        if not texts:
            return np.array([])
        
        try:
            embeddings = self.embedding_model.encode(texts, convert_to_numpy=True)
            return embeddings
        except Exception as e:
            logger.error(f"Failed to generate embeddings: {e}")
            return np.array([])
    
    def get_user_vector_db_path(self, username: str) -> str:
        """Get the path for a user's vector database"""
        return os.path.join(self.vector_db_path, f"{username}.faiss")
    
    def get_user_metadata_path(self, username: str) -> str:
        """Get the path for a user's metadata file"""
        return os.path.join(self.vector_db_path, f"{username}_metadata.pkl")
    
    def add_chunks_to_vector_db(self, username: str, chunks: List[Dict]) -> bool:
        """
        Add text chunks to user's vector database
        
        Args:
            username: User identifier
            chunks: List of chunk dictionaries
            
        Returns:
            Success boolean
        """
        if not chunks:
            return True
        
        try:
            texts = [chunk['text'] for chunk in chunks]
            
            embeddings = self.generate_embeddings(texts)
            if embeddings.size == 0:
                return False
            
            faiss_index_path = self.get_user_vector_db_path(username)
            metadata_path = self.get_user_metadata_path(username)
            
            if os.path.exists(faiss_index_path):
                index = faiss.read_index(faiss_index_path)
                with open(metadata_path, 'rb') as f:
                    existing_metadata = pickle.load(f)
            else:
                index = faiss.IndexFlatIP(self.embedding_dimension)  # Inner product (cosine similarity)
                existing_metadata = []
            
            # Normalize embeddings for cosine similarity
            embeddings = embeddings / np.linalg.norm(embeddings, axis=1, keepdims=True)
            index.add(embeddings)
            
            existing_metadata.extend(chunks)
            
            faiss.write_index(index, faiss_index_path)
            with open(metadata_path, 'wb') as f:
                pickle.dump(existing_metadata, f)
            
            logger.info(f"Added {len(chunks)} chunks to {username}'s vector DB")
            return True
            
        except Exception as e:
            logger.error(f"Failed to add chunks to vector DB: {e}")
            return False
    
    def search_user_vector_db(self, username: str, query: str, top_k: int = 5) -> List[Dict]:
        """
        Search user's vector database for relevant chunks
        
        Args:
            username: User identifier
            query: Search query
            top_k: Number of results to return
            
        Returns:
            List of relevant chunks with scores
        """
        faiss_index_path = self.get_user_vector_db_path(username)
        metadata_path = self.get_user_metadata_path(username)
        
        if not os.path.exists(faiss_index_path) or not os.path.exists(metadata_path):
            return []
        
        try:
            index = faiss.read_index(faiss_index_path)
            with open(metadata_path, 'rb') as f:
                metadata = pickle.load(f)
            
            query_embedding = self.generate_embeddings([query])
            if query_embedding.size == 0:
                return []
            
            query_embedding = query_embedding / np.linalg.norm(query_embedding, axis=1, keepdims=True)
            
            scores, indices = index.search(query_embedding, min(top_k, index.ntotal))
            
            results = []
            for score, idx in zip(scores[0], indices[0]):
                if idx >= 0 and idx < len(metadata):
                    result = {
                        'chunk': metadata[idx],
                        'score': float(score)
                    }
                    results.append(result)
            
            return results
            
        except Exception as e:
            logger.error(f"Failed to search vector DB: {e}")
            return []
    
    def process_file_for_rag(self, file_content: bytes, filename: str, username: str, cid: str) -> bool:
        """
        Complete pipeline to process a file for RAG
        
        Args:
            file_content: Raw file bytes
            filename: Name of the file
            username: User who uploaded the file
            cid: IPFS CID of the file
            
        Returns:
            Success boolean
        """
        try:
            text = self.extract_text_from_file(file_content, filename)
            if not text.strip():
                logger.info(f"No text extracted from {filename}")
                return True
            
            metadata = {
                'username': username,
                'filename': filename,
                'cid': cid,
                'file_type': filename.lower().split('.')[-1]
            }
            
            chunks = self.chunk_text(text, metadata)
            if not chunks:
                logger.info(f"No chunks created from {filename}")
                return True
            
            success = self.add_chunks_to_vector_db(username, chunks)
            
            if success:
                logger.info(f"Successfully processed {filename} for RAG: {len(chunks)} chunks")
            
            return success
            
        except Exception as e:
            logger.error(f"Failed to process file for RAG: {e}")
            return False
    
    def get_user_stats(self, username: str) -> Dict:
        """Get statistics about user's vector database"""
        metadata_path = self.get_user_metadata_path(username)
        
        stats = {
            'total_chunks': 0,
            'total_files': 0,
            'files': []
        }
        
        if os.path.exists(metadata_path):
            try:
                with open(metadata_path, 'rb') as f:
                    metadata = pickle.load(f)
                
                stats['total_chunks'] = len(metadata)
                
                files = set()
                for chunk in metadata:
                    files.add(chunk['metadata']['filename'])
                
                stats['total_files'] = len(files)
                stats['files'] = list(files)
                
            except Exception as e:
                logger.error(f"Failed to get user stats: {e}")
        
        return stats


class LLMIntegration:
    """
    Handles integration with Language Models for generating responses
    """
    
    def __init__(self, model_type: str = "gemini", api_key: str = None):
        """
        Initialize LLM integration
        
        Args:
            model_type: Type of model ("gemini" or "local")
            api_key: API key for external services
        """
        self.model_type = model_type
        
        if model_type == "gemini" and api_key:
            self.api_key = api_key
    
    def generate_answer(self, query: str, context_chunks: List[Dict], max_tokens: int = 5000) -> str:
        """
        Generate an answer using retrieved context
        
        Args:
            query: User's question
            context_chunks: Retrieved relevant chunks
            max_tokens: Maximum tokens for response
            
        Returns:
            Generated answer
        """
        if not context_chunks:
            return "I couldn't find any relevant information in your uploaded files to answer this question."
        
        context_text = "\n\n".join([
            f"From {chunk['chunk']['metadata']['filename']}:\n{chunk['chunk']['text']}"
            for chunk in context_chunks[:3]
        ])
        
        prompt = f"""Based on the following context from the user's uploaded files, please answer their question. If the context doesn't contain enough information to answer the question, please say so.

Context:
{context_text}

Question: {query}

Answer:"""
        
        try:
            if self.model_type == "gemini":
                return self._generate_gemini_response(prompt, max_tokens)
            else:
                return self._generate_simple_response(query, context_chunks)
        except Exception as e:
            logger.error(f"Failed to generate LLM response: {e}")
            return "I encountered an error while generating the response. Please try again."
    
    def _generate_gemini_response(self, prompt: str, max_tokens: int) -> str:
        """Generate response using Gemini API"""
        try:
            model_name = "gemini-2.5-flash"
            url = f"https://generativelanguage.googleapis.com/v1beta/models/{model_name}:generateContent"

            headers = {
                "x-goog-api-key": self.api_key,
                "Content-Type": "application/json"
            }

            data = {
                "contents": [
                    {
                        "parts": [
                            {
                                "text": prompt
                            }
                        ]
                    }
                ],
                "generationConfig": {
                    "maxOutputTokens": max_tokens
                },
                "systemInstruction": {
                    "parts": [
                        {
                            "text": "You are a helpful assistant that answers questions based on the provided context from user's documents."
                        }
                    ]
                }
            }

            response = requests.post(url, headers=headers, json=data)
            response.raise_for_status()

            result = response.json()
            logger.info(f"Gemini API response: {result}")
            return result['candidates'][0]['content']['parts'][0]['text']
        except Exception as e:
            logger.error(f"Gemini API error: {e}")
            raise
    
    def _generate_simple_response(self, query: str, context_chunks: List[Dict]) -> str:
        """Generate a simple extractive response when no LLM is available"""
        if not context_chunks:
            return "No relevant information found."
        
        best_chunk = context_chunks[0]
        filename = best_chunk['chunk']['metadata']['filename']
        text = best_chunk['chunk']['text']
        
        return f"Based on your file '{filename}', here's the most relevant information I found:\n\n{text}..."


rag_manager = None
llm_integration = None

def get_rag_manager() -> RAGManager:
    """Get global RAG manager instance"""
    global rag_manager
    if rag_manager is None:
        rag_manager = RAGManager()
    return rag_manager

def get_llm_integration() -> LLMIntegration:
    """Get global LLM integration instance"""
    global llm_integration
    if llm_integration is None:
        api_key = os.getenv('GOOGLE_API_KEY')
        if api_key:
            llm_integration = LLMIntegration("gemini", api_key)
        else:
            llm_integration = LLMIntegration("simple")
    return llm_integration 